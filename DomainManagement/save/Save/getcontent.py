from bs4 import BeautifulSoup
import requests
import json
from docx import Document

name ="Geturl"


resp = requests.post("https://os.3i.com.vn/PythonCrawler/GetCrawlerData?spiderName=crawler")
resp_json = resp.json()
print(resp_json)

url = resp_json['Url']

# list_child_url = []
content = []

class url_obj:
    def __init__(self, url, iscan, deep):
        self.url = url
        self.iscan = iscan
        self.deep = deep
    Idx = 0
    chk = 0
    with open('C:\\Users\\pycha\\PycharmProjects\\Source\DomainManagement\\DomainManagement\\spiders\\selector5.json', 'r') as j:
            list_tag = json.loads(j.read())

    def GetContent(self, url):
        list_tag = self.list_tag

        req = requests.get('https://en.vneconomy.vn/')
        html = req.text
        soup = BeautifulSoup(html, 'html.parser')
        LstLink = soup.find('html')
        for s in LstLink.find_all(list_tag['title']):
            try:
                title = s.text
                if '' != title:
                    content.append(title)
            except:
                pass
        for s in LstLink.find_all(list_tag['text']):
            try:
                text =s.text
                if '' != text:
                    content.append(text)
            except:
                pass

        # time = LstLink.find_all(list_tag['time'], class_= list_tag['time_class'])
        # link_image = s.findAll(p['img']
        # link_image = img[p['img']]
        # img_head = requests.get(link_image)
        for img in LstLink.find_all(list_tag['img'], class_= list_tag['img_class']):
            image = (img.get('src'))
            print(image)

                    # with open("link_image", "w+b") as im_h:
                    #     im_h.write(link_image.content)
        # if '' != title:
        #     content.append({
        #         title,
        #         text,
        #         # time,                  # image
        #         })
        return content

# list_child_url.append(url_obj(url, 0, 1))
object = url_obj(url, 0, 1)
# object.main()
document = Document()
object.GetContent(url)
# for obj in list_child_url:
#     document.add_paragraph(obj.url)
#     print(obj.url, obj.iscan, obj.deep, sep=' ')
# print(len(list_child_url))
for value in content:
    document.add_paragraph(value)
    #document.add_picture('img_head.png')
    print(value)
print(len(content))
# Vị trí lưu file
filePath = '/DomainManagement\DomainManagement\spiders/url1.docx'
document.save("C:\\Users\\pycha\\PycharmProjects\\Source" + filePath)